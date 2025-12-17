
from docx import Document
from app.processors.document_processor import DocumentProcessor, ProcessingError
from app.models.data_models import ProcessedDocument


class DocProcessor(DocumentProcessor):
    
    def can_process(self, file_path: str, file_type: str) -> bool:
        doc_types = ['.doc', '.docx']
        mime_types = ['application/msword', 
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        return file_type.lower() in doc_types or file_type in mime_types
    
    def process_document(self, file_path: str, filename: str) -> ProcessedDocument:
        try:
            file_info = self._get_file_info(file_path, filename)
            text_content = self._extract_text_from_doc(file_path)
            
            return ProcessedDocument(
                filename=filename,
                file_type=file_info["file_extension"],
                content=text_content,
                file_size=file_info["file_size_bytes"],
                metadata={"processor": "DocProcessor"}
            )
            
        except Exception as e:
            raise ProcessingError(f"Failed to process DOC/DOCX {filename}: {str(e)}")
    
    def _extract_text_from_doc(self, file_path: str) -> str:
        if file_path.lower().endswith('.docx'):
            return self._extract_from_docx(file_path)
        else:
            return self._extract_from_legacy_doc(file_path)
    
    def _extract_from_docx(self, file_path: str) -> str:
        doc = None
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if text_parts:
                return '\n\n'.join(text_parts)
            else:
                return "No readable text content found in this document."
                
        except Exception as e:
            error_msg = str(e).lower()
            if 'word/null' in error_msg or 'archive' in error_msg or 'zipfile' in error_msg:
                raise ProcessingError(f"The DOCX file appears to be corrupted or invalid. Please try re-saving the document or use a different file format.")
            else:
                raise ProcessingError(f"DOCX processing failed: {str(e)}")
        finally:
            if doc is not None:
                try:
                    if hasattr(doc, 'part') and hasattr(doc.part, 'package'):
                        doc.part.package.close()
                except:
                    pass
    
    def _extract_from_legacy_doc(self, file_path: str) -> str:
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
        except ImportError:
            pass
        
        return "Legacy DOC format detected. Please convert to DOCX for better text extraction."